from __future__ import annotations

import base64
from collections import Counter, defaultdict
from datetime import datetime
from io import BytesIO
from statistics import mean, median
from typing import Any, Dict, Iterable, List, Optional

from rdkit import Chem
from rdkit.Chem import Draw
from rdkit.Chem.Scaffolds import MurckoScaffold
from PIL import Image, ImageDraw, ImageFont

from app.services.sar_summary import generate_sar_summary


DESCRIPTOR_COLUMNS = [
    ("mol_weight", "Molecular weight"),
    ("logp", "cLogP"),
    ("tpsa", "TPSA"),
    ("hbd", "HBD"),
    ("hba", "HBA"),
    ("rotatable_bonds", "Rotatable bonds"),
]


def build_briefing_report(
    molecules: List[Dict[str, Any]],
    project_name: str,
    potency_column: Optional[str] = None,
    potency_direction: str = "lower_is_better",
    admet_columns: Optional[List[str]] = None,
    min_fold_change: float = 3.0,
) -> Dict[str, Any]:
    admet_columns = admet_columns or []
    sar_summary = _safe_sar_summary(molecules, potency_column, potency_direction, admet_columns, min_fold_change)
    dataset_summary = _dataset_summary(molecules, potency_column)
    property_distribution = _property_distribution(molecules, potency_column, admet_columns)
    scaffold_clusters = _scaffold_clusters(molecules, potency_column)
    limitations = _limitations(molecules, sar_summary, potency_column)

    report = {
        "title": project_name,
        "generated_at": datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC"),
        "molecule_lookup": {int(molecule["id"]): molecule for molecule in molecules},
        "dataset_summary": dataset_summary,
        "property_distribution": property_distribution,
        "scaffold_clusters": scaffold_clusters,
        "sar_summary": sar_summary,
        "limitations": limitations,
    }
    report["publication_figure_png"] = _publication_figure_png(report)
    report["markdown"] = _render_markdown(report)
    return report


def build_briefing_docx(report: Dict[str, Any]) -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[style_name].font.name = "Arial"
        styles[style_name].font.color.rgb = RGBColor(20, 35, 45)

    title = document.add_heading(report["title"], level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    subtitle = document.add_paragraph(f"Medicinal chemistry briefing report | Generated {report['generated_at']}")
    subtitle.runs[0].font.color.rgb = RGBColor(79, 97, 110)

    _add_heading(document, "Project Overview")
    _add_paragraph(
        document,
        "This briefing summarizes uploaded compound structures, calculated physicochemical properties, "
        "Bemis-Murcko scaffold clusters, and available SAR evidence. It is intended as a project-ready "
        "starting point for medicinal chemistry review.",
    )

    _add_heading(document, "Dataset Summary")
    _add_key_value_table(document, report["dataset_summary"])

    _add_heading(document, "Property Distribution")
    _add_distribution_table(document, report["property_distribution"])

    _add_heading(document, "Scaffold Clusters")
    _add_scaffold_table(document, report["scaffold_clusters"])

    _add_heading(document, "SAR Observations")
    sar = report["sar_summary"]
    for item in sar.get("key_sar_trends", []):
        document.add_paragraph(item, style="List Bullet")

    if report.get("publication_figure_png"):
        _add_heading(document, "Publication SAR Figure")
        document.add_picture(BytesIO(report["publication_figure_png"]), width=Inches(6.8))

    _add_heading(document, "Recommended Next Analogs")
    for item in sar.get("suggested_next_analogs", []):
        document.add_paragraph(item, style="List Bullet")

    _add_heading(document, "Limitations")
    for item in report["limitations"]:
        document.add_paragraph(item, style="List Bullet")

    if sar.get("evidence_table"):
        _add_heading(document, "SAR Evidence Table")
        _add_evidence_table(document, sar["evidence_table"][:12], report["molecule_lookup"])

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _safe_sar_summary(
    molecules: List[Dict[str, Any]],
    potency_column: Optional[str],
    potency_direction: str,
    admet_columns: List[str],
    min_fold_change: float,
) -> Dict[str, Any]:
    if not potency_column:
        return {
            "key_sar_trends": ["No potency column was supplied, so SAR trend detection was skipped."],
            "risky_modifications": ["No ADMET-linked risky modifications were evaluated."],
            "promising_positions": ["No potency-linked promising positions were evaluated."],
            "suggested_next_analogs": ["Upload potency values to generate matched-pair next-analog suggestions."],
            "evidence_table": [],
            "matched_pair_count": 0,
            "scaffold_count": 0,
        }
    try:
        return generate_sar_summary(molecules, potency_column, potency_direction, admet_columns, min_fold_change)
    except ValueError as exc:
        return {
            "key_sar_trends": [f"SAR trend detection could not run: {exc}"],
            "risky_modifications": ["No risky modifications were evaluated."],
            "promising_positions": ["No promising positions were evaluated."],
            "suggested_next_analogs": ["Check potency settings and rerun the report."],
            "evidence_table": [],
            "matched_pair_count": 0,
            "scaffold_count": 0,
        }


def _dataset_summary(molecules: List[Dict[str, Any]], potency_column: Optional[str]) -> Dict[str, str]:
    sources = sorted({molecule.get("source_filename", "unknown") for molecule in molecules})
    potency_count = sum(1 for molecule in molecules if potency_column and _parse_number(_property_value(molecule, potency_column)) is not None)
    return {
        "Compounds": str(len(molecules)),
        "Source files": ", ".join(sources) if sources else "None",
        "Compounds with potency": str(potency_count) if potency_column else "Not evaluated",
        "Potency column": potency_column or "Not supplied",
    }


def _property_distribution(
    molecules: List[Dict[str, Any]],
    potency_column: Optional[str],
    admet_columns: List[str],
) -> List[Dict[str, str]]:
    distributions = []
    descriptors = DESCRIPTOR_COLUMNS[:]
    if potency_column:
        descriptors.append((potency_column, potency_column))
    for column in admet_columns:
        descriptors.append((column, column))

    for column, label in descriptors:
        values = [_numeric_value(molecule, column) for molecule in molecules]
        values = [value for value in values if value is not None]
        if not values:
            continue
        distributions.append(
            {
                "Property": label,
                "N": str(len(values)),
                "Min": _format_number(min(values)),
                "Median": _format_number(median(values)),
                "Mean": _format_number(mean(values)),
                "Max": _format_number(max(values)),
            }
        )
    return distributions


def _scaffold_clusters(molecules: List[Dict[str, Any]], potency_column: Optional[str]) -> List[Dict[str, str]]:
    groups: Dict[str, List[Dict[str, Any]]] = defaultdict(list)
    for molecule in molecules:
        scaffold = _murcko_scaffold(molecule.get("smiles", ""))
        groups[scaffold].append(molecule)

    rows = []
    for scaffold, members in groups.items():
        potencies = [_numeric_value(member, potency_column) for member in members] if potency_column else []
        potencies = [value for value in potencies if value is not None]
        rows.append(
            {
                "Scaffold": scaffold,
                "Compounds": str(len(members)),
                "Representative": members[0].get("name") or str(members[0].get("id")),
                "Median potency": _format_number(median(potencies)) if potencies else "N/A",
                "Avg MW": _format_number(_average(member.get("mol_weight") for member in members)),
                "Avg cLogP": _format_number(_average(member.get("logp") for member in members)),
            }
        )
    return sorted(rows, key=lambda row: int(row["Compounds"]), reverse=True)[:12]


def _limitations(
    molecules: List[Dict[str, Any]],
    sar_summary: Dict[str, Any],
    potency_column: Optional[str],
) -> List[str]:
    limitations = [
        "PDF extraction and assay-table normalization are not included in this MVP report path.",
        "Scaffold clusters use Bemis-Murcko scaffolds and may overgroup or undergroup chemotypes depending on core definition.",
        "Matched molecular pairs are heuristic and should be reviewed by a chemist before committing synthesis priorities.",
    ]
    if not potency_column:
        limitations.append("No potency column was supplied, so potency-ranked SAR conclusions are unavailable.")
    elif sar_summary.get("matched_pair_count", 0) == 0:
        limitations.append("No matched pairs met the fold-change threshold; recommendations are limited by sparse SAR evidence.")
    if len(molecules) < 10:
        limitations.append("The dataset is small, so property distributions and trend confidence are directional.")
    return limitations


def _render_markdown(report: Dict[str, Any]) -> str:
    sar = report["sar_summary"]
    lines = [
        f"# {report['title']}",
        "",
        f"Generated: {report['generated_at']}",
        "",
        "## Project Overview",
        "",
        "This medicinal chemistry briefing summarizes the uploaded compound dataset, calculated properties, scaffold clusters, SAR observations, and recommended next analogs.",
        "",
        "## Dataset Summary",
        "",
        _markdown_kv_table(report["dataset_summary"]),
        "",
        "## Property Distribution",
        "",
        _markdown_table(report["property_distribution"]) or "No numeric property distributions were available.",
        "",
        "## Scaffold Clusters",
        "",
        _markdown_table(report["scaffold_clusters"]) or "No valid scaffold clusters were generated.",
        "",
        "## SAR Observations",
        "",
        *[f"- {item}" for item in sar.get("key_sar_trends", [])],
    ]
    if report.get("publication_figure_png"):
        lines.extend(
            [
                "",
                "## Publication SAR Figure",
                "",
                f"![Publication-quality SAR summary plate]({_png_data_uri(report['publication_figure_png'])})",
            ]
        )
    lines.extend(
        [
            "",
            "## Recommended Next Analogs",
            "",
            *[f"- {item}" for item in sar.get("suggested_next_analogs", [])],
            "",
            "## Limitations",
            "",
            *[f"- {item}" for item in report["limitations"]],
        ]
    )
    if sar.get("evidence_table"):
        lines.extend(
            [
                "",
                "## SAR Evidence Table",
                "",
                _markdown_table(_evidence_rows_for_markdown(sar["evidence_table"], report["molecule_lookup"], include_images=True)),
            ]
        )
    return "\n".join(lines).strip() + "\n"


def _markdown_kv_table(items: Dict[str, str]) -> str:
    return _markdown_table([{"Metric": key, "Value": value} for key, value in items.items()])


def _markdown_table(rows: List[Dict[str, Any]]) -> str:
    if not rows:
        return ""
    columns = list(rows[0].keys())
    output = [
        "| " + " | ".join(columns) + " |",
        "| " + " | ".join("---" for _ in columns) + " |",
    ]
    for row in rows:
        output.append("| " + " | ".join(_escape_markdown(row.get(column, "")) for column in columns) + " |")
    return "\n".join(output)


def _evidence_rows_for_markdown(
    rows: List[Dict[str, Any]],
    molecule_lookup: Optional[Dict[int, Dict[str, Any]]] = None,
    include_images: bool = False,
) -> List[Dict[str, Any]]:
    return [
        {
            "Pair": f"{row.get('weaker_name') or row['weaker_compound_id']} -> {row.get('stronger_name') or row['stronger_compound_id']}",
            **_markdown_structure_columns(row, molecule_lookup or {}, include_images),
            "Position": row.get("position") or "multi-site",
            "Change": f"{row['weaker_substitution']} -> {row['stronger_substitution']}",
            "Fold": f"{row['fold_change']:.1f}x",
            "Note": row["note"],
        }
        for row in rows[:25]
    ]


def _add_heading(document: Any, text: str) -> None:
    document.add_heading(text, level=1)


def _add_paragraph(document: Any, text: str) -> None:
    document.add_paragraph(text)


def _add_key_value_table(document: Any, items: Dict[str, str]) -> None:
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Metric"
    table.rows[0].cells[1].text = "Value"
    for key, value in items.items():
        cells = table.add_row().cells
        cells[0].text = key
        cells[1].text = value


def _add_distribution_table(document: Any, rows: List[Dict[str, str]]) -> None:
    _add_rows_table(document, rows, ["Property", "N", "Min", "Median", "Mean", "Max"])


def _add_scaffold_table(document: Any, rows: List[Dict[str, str]]) -> None:
    _add_rows_table(document, rows, ["Scaffold", "Compounds", "Representative", "Median potency", "Avg MW", "Avg cLogP"])


def _add_evidence_table(document: Any, rows: List[Dict[str, Any]], molecule_lookup: Dict[int, Dict[str, Any]]) -> None:
    from docx.shared import Inches

    if not rows:
        document.add_paragraph("No SAR evidence available.")
        return

    columns = ["Pair", "Weaker", "Stronger", "Position", "Change", "Fold", "Note"]
    table = document.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    for index, column in enumerate(columns):
        table.rows[0].cells[index].text = column

    for row in rows:
        cells = table.add_row().cells
        cells[0].text = f"{row.get('weaker_name') or row['weaker_compound_id']} -> {row.get('stronger_name') or row['stronger_compound_id']}"
        _add_structure_to_cell(cells[1], _smiles_for_id(molecule_lookup, row["weaker_compound_id"]), Inches(1.05))
        _add_structure_to_cell(cells[2], _smiles_for_id(molecule_lookup, row["stronger_compound_id"]), Inches(1.05))
        cells[3].text = row.get("position") or "multi-site"
        cells[4].text = f"{row['weaker_substitution']} -> {row['stronger_substitution']}"
        cells[5].text = f"{row['fold_change']:.1f}x"
        cells[6].text = row["note"]


def _add_rows_table(document: Any, rows: List[Dict[str, Any]], columns: List[str]) -> None:
    if not rows:
        document.add_paragraph("No data available.")
        return
    table = document.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    for index, column in enumerate(columns):
        table.rows[0].cells[index].text = column
    for row in rows:
        cells = table.add_row().cells
        for index, column in enumerate(columns):
            cells[index].text = str(row.get(column, ""))


def _murcko_scaffold(smiles: str) -> str:
    mol = Chem.MolFromSmiles(smiles)
    if mol is None:
        return "Invalid structure"
    scaffold = MurckoScaffold.GetScaffoldForMol(mol)
    if scaffold.GetNumAtoms() == 0:
        return Chem.MolToSmiles(mol, canonical=True)
    return Chem.MolToSmiles(scaffold, canonical=True)


def _markdown_structure_columns(
    row: Dict[str, Any],
    molecule_lookup: Dict[int, Dict[str, Any]],
    include_images: bool,
) -> Dict[str, str]:
    if not include_images:
        return {}
    weaker_smiles = _smiles_for_id(molecule_lookup, row["weaker_compound_id"])
    stronger_smiles = _smiles_for_id(molecule_lookup, row["stronger_compound_id"])
    return {
        "Weaker Structure": _markdown_structure_image(weaker_smiles, row.get("weaker_name") or str(row["weaker_compound_id"])),
        "Stronger Structure": _markdown_structure_image(stronger_smiles, row.get("stronger_name") or str(row["stronger_compound_id"])),
    }


def _smiles_for_id(molecule_lookup: Dict[int, Dict[str, Any]], molecule_id: int) -> Optional[str]:
    molecule = molecule_lookup.get(int(molecule_id))
    return molecule.get("smiles") if molecule else None


def _markdown_structure_image(smiles: Optional[str], label: str) -> str:
    data_uri = _structure_svg_data_uri(smiles)
    if not data_uri:
        return "N/A"
    return f"![{_escape_markdown(label)}]({data_uri})"


def _structure_svg_data_uri(smiles: Optional[str]) -> Optional[str]:
    if not smiles:
        return None
    mol = Chem.MolFromSmiles(smiles)
    if mol is None:
        return None
    drawer = Draw.MolDraw2DSVG(220, 150)
    drawer.DrawMolecule(mol)
    drawer.FinishDrawing()
    svg = drawer.GetDrawingText()
    encoded = base64.b64encode(svg.encode("utf-8")).decode("ascii")
    return f"data:image/svg+xml;base64,{encoded}"


def _structure_png_bytes(smiles: Optional[str], size: tuple[int, int] = (320, 220)) -> Optional[BytesIO]:
    if not smiles:
        return None
    mol = Chem.MolFromSmiles(smiles)
    if mol is None:
        return None
    image = Draw.MolToImage(mol, size=size)
    buffer = BytesIO()
    image.save(buffer, format="PNG")
    buffer.seek(0)
    return buffer


def _add_structure_to_cell(cell: Any, smiles: Optional[str], width: Any) -> None:
    image = _structure_png_bytes(smiles)
    if image is None:
        cell.text = "N/A"
        return
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.add_run().add_picture(image, width=width)


def _publication_figure_png(report: Dict[str, Any]) -> Optional[bytes]:
    evidence = report["sar_summary"].get("evidence_table", [])
    if not evidence:
        return None

    molecule_lookup = report["molecule_lookup"]
    width = 2400
    table_rows = min(len(evidence), 8)
    row_height = 132
    height = 620 + table_rows * row_height
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)

    title_font = _font("Times New Roman Italic.ttf", 40, italic=True)
    caption_font = _font("Times New Roman.ttf", 34)
    header_font = _font("Times New Roman Bold.ttf", 28, bold=True)
    body_font = _font("Times New Roman.ttf", 27)
    small_font = _font("Times New Roman.ttf", 23)

    draw.text((55, 38), _fit_text(report["title"], title_font, 980), fill=(20, 20, 20), font=title_font)
    draw.text((1600, 38), "Medicinal chemistry briefing", fill=(20, 20, 20), font=caption_font)
    draw.line((55, 92, width - 55, 92), fill=(20, 20, 20), width=3)
    draw.text((55, 132), "SAR Figure. Matched-pair substitutions associated with potency shifts.", fill=(20, 20, 20), font=header_font)

    top_pairs = evidence[:3]
    panel_top = 188
    panel_width = (width - 140) // max(len(top_pairs), 1)
    arrow_font = _font("Arial.ttf", 46)
    for index, row in enumerate(top_pairs):
        x = 70 + index * panel_width
        pair_label = f"{row.get('weaker_name') or row['weaker_compound_id']} -> {row.get('stronger_name') or row['stronger_compound_id']}"
        draw.text((x + 16, panel_top), pair_label, fill=(20, 20, 20), font=header_font)
        draw.text((x + 16, panel_top + 36), f"{row['fold_change']:.1f}x improvement; {row.get('position') or 'multi-site'}", fill=(70, 70, 70), font=small_font)
        weaker_image = _mol_image(_smiles_for_id(molecule_lookup, row["weaker_compound_id"]), (310, 210))
        stronger_image = _mol_image(_smiles_for_id(molecule_lookup, row["stronger_compound_id"]), (310, 210))
        if weaker_image:
            image.paste(weaker_image, (x + 10, panel_top + 74))
        draw.text((x + 330, panel_top + 152), "→", fill=(25, 25, 25), font=arrow_font)
        if stronger_image:
            image.paste(stronger_image, (x + 400, panel_top + 74))
        draw.text((x + 26, panel_top + 298), row["weaker_substitution"], fill=(20, 20, 20), font=small_font)
        draw.text((x + 420, panel_top + 298), row["stronger_substitution"], fill=(20, 20, 20), font=small_font)

    table_top = 560
    draw.line((55, table_top - 24, width - 55, table_top - 24), fill=(20, 20, 20), width=3)
    columns = [
        ("cmpd", 55, 190),
        ("weaker", 190, 470),
        ("stronger", 470, 750),
        ("position", 750, 930),
        ("substitution change", 930, 1420),
        ("fold", 1420, 1535),
        ("SAR note", 1535, width - 55),
    ]
    for label, x0, _ in columns:
        draw.text((x0 + 8, table_top), label, fill=(20, 20, 20), font=header_font)
    draw.line((55, table_top + 42, width - 55, table_top + 42), fill=(20, 20, 20), width=2)

    for index, row in enumerate(evidence[:table_rows]):
        y = table_top + 58 + index * row_height
        if index % 2:
            draw.rectangle((55, y - 8, width - 55, y + row_height - 12), fill=(248, 248, 248))
        values = [
            str(index + 1),
            str(row.get("weaker_name") or row["weaker_compound_id"]),
            str(row.get("stronger_name") or row["stronger_compound_id"]),
            str(row.get("position") or "multi-site"),
            f"{row['weaker_substitution']} -> {row['stronger_substitution']}",
            f"{row['fold_change']:.1f}x",
            row["note"],
        ]
        for value, (_, x0, x1) in zip(values, columns):
            _draw_wrapped(draw, value, x0 + 8, y, x1 - x0 - 16, body_font, fill=(20, 20, 20), max_lines=3)
        draw.line((55, y + row_height - 12, width - 55, y + row_height - 12), fill=(210, 210, 210), width=1)

    draw.text(
        (55, height - 42),
        "Structures generated from uploaded SMILES with RDKit; matched pairs are heuristic and require medicinal chemistry review.",
        fill=(80, 80, 80),
        font=small_font,
    )
    buffer = BytesIO()
    image.save(buffer, format="PNG", dpi=(300, 300))
    return buffer.getvalue()


def _mol_image(smiles: Optional[str], size: tuple[int, int]) -> Optional[Image.Image]:
    if not smiles:
        return None
    mol = Chem.MolFromSmiles(smiles)
    if mol is None:
        return None
    return Draw.MolToImage(mol, size=size).convert("RGB")


def _png_data_uri(png_bytes: bytes) -> str:
    encoded = base64.b64encode(png_bytes).decode("ascii")
    return f"data:image/png;base64,{encoded}"


def _font(name: str, size: int, bold: bool = False, italic: bool = False) -> ImageFont.ImageFont:
    candidates = [f"/System/Library/Fonts/Supplemental/{name}"]
    if bold:
        candidates.insert(0, "/System/Library/Fonts/Supplemental/Times New Roman Bold.ttf")
    if italic:
        candidates.insert(0, "/System/Library/Fonts/Supplemental/Times New Roman Italic.ttf")
    candidates.extend(
        [
            "/System/Library/Fonts/Supplemental/Times New Roman.ttf",
            "/System/Library/Fonts/Supplemental/Arial.ttf",
            "/System/Library/Fonts/Helvetica.ttc",
        ]
    )
    for path in candidates:
        try:
            return ImageFont.truetype(path, size=size)
        except OSError:
            continue
    return ImageFont.load_default()


def _fit_text(text: str, font: ImageFont.ImageFont, max_width: int) -> str:
    if font.getlength(text) <= max_width:
        return text
    trimmed = text
    while trimmed and font.getlength(f"{trimmed}...") > max_width:
        trimmed = trimmed[:-1]
    return f"{trimmed}..."


def _draw_wrapped(
    draw: ImageDraw.ImageDraw,
    text: str,
    x: int,
    y: int,
    max_width: int,
    font: ImageFont.ImageFont,
    fill: tuple[int, int, int],
    max_lines: int,
) -> None:
    words = str(text).split()
    lines: List[str] = []
    current = ""
    for word in words:
        candidate = f"{current} {word}".strip()
        if draw.textlength(candidate, font=font) <= max_width:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    if len(lines) > max_lines:
        lines = lines[:max_lines]
        lines[-1] = _fit_text(lines[-1], font, max_width)
    for index, line in enumerate(lines):
        draw.text((x, y + index * 31), line, fill=fill, font=font)


def _property_value(molecule: Dict[str, Any], column: str) -> Any:
    if column in molecule:
        return molecule[column]
    properties = molecule.get("properties", {})
    if column in properties:
        return properties[column]
    lowered = {str(key).lower(): value for key, value in properties.items()}
    return lowered.get(column.lower())


def _numeric_value(molecule: Dict[str, Any], column: Optional[str]) -> Optional[float]:
    if not column:
        return None
    return _parse_number(_property_value(molecule, column))


def _parse_number(value: Any) -> Optional[float]:
    if value is None:
        return None
    if isinstance(value, (int, float)):
        return float(value)
    text = str(value).replace(",", "")
    number = ""
    for char in text:
        if char.isdigit() or char in ".-+eE":
            number += char
        elif number:
            break
    try:
        return float(number) if number else None
    except ValueError:
        return None


def _present_values(values: Iterable[Any]) -> List[float]:
    return [float(value) for value in values if value is not None]


def _average(values: Iterable[Any]) -> Optional[float]:
    present = _present_values(values)
    return mean(present) if present else None


def _format_number(value: Any) -> str:
    if isinstance(value, list):
        if not value:
            return "N/A"
        value = mean(value)
    try:
        number = float(value)
    except (TypeError, ValueError):
        return "N/A"
    return f"{number:.2f}"


def _escape_markdown(value: Any) -> str:
    return str(value).replace("|", "\\|").replace("\n", " ")
