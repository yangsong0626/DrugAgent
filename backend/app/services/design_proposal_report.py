from __future__ import annotations

from datetime import datetime
from io import BytesIO
from typing import Any

from app.services.next_round_design import generate_next_round_designs
from app.services.sar_workbench import build_sar_workbench


def build_design_proposal_report(
    molecules: list[dict[str, Any]],
    project_name: str,
    potency_column: str | None = None,
    potency_direction: str = "lower_is_better",
    admet_columns: list[str] | None = None,
    objectives: dict[str, Any] | None = None,
    constraints: dict[str, Any] | None = None,
    count: int = 24,
    min_fold_change: float = 3.0,
    preference_profile: dict[str, Any] | None = None,
) -> dict[str, Any]:
    admet_columns = admet_columns or []
    design = generate_next_round_designs(
        molecules=molecules,
        potency_column=potency_column,
        potency_direction=potency_direction,
        admet_columns=admet_columns,
        objectives=objectives,
        constraints=constraints,
        count=count,
        preference_profile=preference_profile,
    )
    resolved_potency = design.get("potency_column")
    sar = _safe_workbench(molecules, resolved_potency, potency_direction, design.get("admet_columns", []), min_fold_change)
    report = {
        "title": project_name,
        "generated_at": datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC"),
        "compound_count": len(molecules),
        "design": design,
        "sar": sar,
    }
    report["markdown"] = _render_design_markdown(report)
    return report


def build_design_proposal_docx(report: dict[str, Any]) -> bytes:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[style_name].font.name = "Arial"
        styles[style_name].font.color.rgb = RGBColor(20, 35, 45)

    document.add_heading(report["title"], level=0)
    document.add_paragraph(f"Next-round medicinal chemistry design proposal | Generated {report['generated_at']}")
    document.add_heading("Executive Summary", level=1)
    for item in _executive_bullets(report):
        document.add_paragraph(item, style="List Bullet")

    document.add_heading("SAR Basis", level=1)
    for item in report["sar"].get("summary", {}).get("key_sar_trends", [])[:6]:
        document.add_paragraph(item, style="List Bullet")

    document.add_heading("Recommended Next Round", level=1)
    _add_recommendation_table(document, report["design"].get("recommendations", [])[:18])

    document.add_heading("Decision Log Entry", level=1)
    document.add_paragraph(_decision_summary(report))

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def decision_log_body_from_report(report: dict[str, Any]) -> dict[str, Any]:
    recommendations = report["design"].get("recommendations", [])
    return {
        "generated_at": report["generated_at"],
        "summary": _decision_summary(report),
        "potency_column": report["design"].get("potency_column"),
        "admet_columns": report["design"].get("admet_columns", []),
        "recommendation_count": len(recommendations),
        "top_recommendations": [
            {
                "name": item["name"],
                "smiles": item["smiles"],
                "score": item["score"],
                "priority": item["priority"],
                "expected_benefit": item["expected_benefit"],
                "main_risk": item["main_risk"],
                "source_molecule_id": item["source_molecule_id"],
            }
            for item in recommendations[:8]
        ],
    }


def _safe_workbench(
    molecules: list[dict[str, Any]],
    potency_column: str | None,
    potency_direction: str,
    admet_columns: list[str],
    min_fold_change: float,
) -> dict[str, Any]:
    if not potency_column:
        return {"summary": {"key_sar_trends": ["No potency endpoint was available for SAR trend generation."]}}
    try:
        return build_sar_workbench(
            molecules=molecules,
            potency_column=potency_column,
            potency_direction=potency_direction,
            admet_columns=admet_columns,
            min_fold_change=min_fold_change,
        )
    except ValueError as exc:
        return {"summary": {"key_sar_trends": [f"SAR workbench could not run: {exc}"]}}


def _render_design_markdown(report: dict[str, Any]) -> str:
    design = report["design"]
    sar = report["sar"].get("summary", {})
    recommendations = design.get("recommendations", [])
    lines = [
        f"# {report['title']}",
        "",
        f"Generated: {report['generated_at']}",
        "",
        "## Executive Summary",
        "",
        *[f"- {item}" for item in _executive_bullets(report)],
        "",
        "## SAR Basis",
        "",
        *[f"- {item}" for item in sar.get("key_sar_trends", [])[:8]],
        "",
        "## Next-Round Design Recommendations",
        "",
        _recommendation_markdown_table(recommendations[:24]) or "No next-round recommendations were generated.",
        "",
        "## Top Recommendation Detail",
        "",
    ]
    for item in recommendations[:8]:
        lines.extend(
            [
                f"### {item['name']} ({item['priority']}, score {item['score']})",
                "",
                f"- SMILES: `{item['smiles']}`",
                f"- Transform: {item.get('transform_title') or 'N/A'}",
                f"- Expected benefit: {item['expected_benefit']}",
                f"- Main risk: {item['main_risk']}",
                f"- Synthesis: {item['synthetic_feasibility']['level']} ({item['synthetic_feasibility']['score']}) - {item['synthetic_note']}",
                f"- Rationale: {item['rationale']}",
                "",
            ]
        )
    lines.extend(
        [
            "## Decision Record",
            "",
            _decision_summary(report),
        ]
    )
    return "\n".join(lines).strip() + "\n"


def _executive_bullets(report: dict[str, Any]) -> list[str]:
    design = report["design"]
    recommendations = design.get("recommendations", [])
    high_priority = sum(1 for item in recommendations if item.get("priority") == "high")
    easy = sum(1 for item in recommendations if item.get("synthetic_feasibility", {}).get("level") == "easy")
    bullets = [
        f"Analyzed {report['compound_count']} uploaded compounds and generated {len(recommendations)} ranked next-round analogs.",
        f"{high_priority} recommendation(s) are high priority; {easy} have easy synthetic feasibility by the current rule set.",
    ]
    if design.get("potency_column"):
        bullets.append(f"Optimization is anchored to {design['potency_column']} with {design.get('potency_direction', 'lower_is_better')}.")
    if design.get("admet_columns"):
        bullets.append(f"ADMET trade-offs tracked: {', '.join(design['admet_columns'][:5])}.")
    return bullets


def _decision_summary(report: dict[str, Any]) -> str:
    recommendations = report["design"].get("recommendations", [])
    if not recommendations:
        return "No make/test recommendation was recorded because the design engine did not generate candidates."
    top = recommendations[0]
    return (
        f"Prioritize {top['name']} and related high-scoring analogs for the next make/test cycle. "
        f"Lead rationale: {top['expected_benefit']}. Main risk to monitor: {top['main_risk']}."
    )


def _recommendation_markdown_table(recommendations: list[dict[str, Any]]) -> str:
    if not recommendations:
        return ""
    rows = [
        {
            "Priority": item["priority"],
            "Score": item["score"],
            "Name": item["name"],
            "Source": item.get("source_molecule_name") or item["source_molecule_id"],
            "Transform": item.get("transform_title") or "N/A",
            "Benefit": item["expected_benefit"],
            "Synthesis": item["synthetic_feasibility"]["level"],
            "Risk": item["main_risk"],
        }
        for item in recommendations
    ]
    return _markdown_table(rows)


def _markdown_table(rows: list[dict[str, Any]]) -> str:
    if not rows:
        return ""
    columns = list(rows[0].keys())
    output = [
        "| " + " | ".join(columns) + " |",
        "| " + " | ".join("---" for _ in columns) + " |",
    ]
    for row in rows:
        output.append("| " + " | ".join(_escape_markdown(row.get(column, "")) for column in columns) + " |")
    return "\n".join(output)


def _add_recommendation_table(document: Any, recommendations: list[dict[str, Any]]) -> None:
    if not recommendations:
        document.add_paragraph("No next-round recommendations were generated.")
        return
    columns = ["Priority", "Score", "Name", "Transform", "Benefit", "Synthesis", "Risk"]
    table = document.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    for index, column in enumerate(columns):
        table.rows[0].cells[index].text = column
    for item in recommendations:
        cells = table.add_row().cells
        values = [
            item["priority"],
            str(item["score"]),
            item["name"],
            item.get("transform_title") or "N/A",
            item["expected_benefit"],
            item["synthetic_feasibility"]["level"],
            item["main_risk"],
        ]
        for index, value in enumerate(values):
            cells[index].text = str(value)


def _escape_markdown(value: Any) -> str:
    return str(value).replace("|", "\\|").replace("\n", " ")
